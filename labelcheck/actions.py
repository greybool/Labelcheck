"""План действий по отчёту (день 9, требование Сергея): короткие понятные
списки «что делать», без цитат регламентов и канцелярита.

Отчёт вердиктов отвечает на вопрос «что нашли и на каком основании» — он
подробный и с цитатами, потому что вердикт обязан быть проверяемым. Но
работать по нему неудобно: закупщику нужно три коротких списка, которые
можно скопировать в письмо:

- ДИЗАЙНЕРУ  — что поправить в макете (где, что, как должно быть);
- ПОСТАВЩИКУ — что запросить/уточнить у производителя;
- ПРОВЕРИТЬ  — что специалист смотрит сам (замеры, графика, документы).

Как формируется: один вызов дешёвой модели (CHEAP_MODEL) на весь отчёт —
на вход идут только вердикты со статусами «возможное нарушение» и
«требует ручной проверки» (что нашли + объяснение), на выход JSON с
короткими пунктами. Вердикты не переспрашиваются и не меняются: это
пересказ уже готовых выводов, поэтому дёшево (~$0.01) и не трогает кэш
вердиктов.

Правки человека учитываются: если специалист отметил вердикт как ошибку
системы (👎), пункт в план не попадает; выбранный человеком адресат
(дизайнер / поставщик / проверить сам) переопределяет предложенный
моделью.

Без API (нет ключа, сбой сети) модуль не падает: fallback собирает план
из последних предложений объяснений — хуже по стилю, но работает.
"""

import hashlib
import json
import os
import re

from labelcheck.retrieval import ROOT, load_config
from labelcheck.verdict import (STATUS_MANUAL, STATUS_VIOLATION,
                                _load_json_cache, _save_json_cache)
from labelcheck.vision import TokenTally

# Адресаты пунктов плана: ключ → (заголовок, что это значит)
TARGETS = {
    "designer": ("Замечания дизайнеру", "поправить в макете"),
    "supplier": ("Запросить у поставщика", "уточнить у производителя"),
    "manual": ("Проверить самостоятельно", "сверить вручную"),
}
TARGET_KEYS = tuple(TARGETS)

SYSTEM_PROMPT = """Ты — помощник специалиста по закупкам. Тебе дают список
замечаний автоматической проверки макета упаковки. Преврати их в короткий
план работ на русском языке.

Верни JSON:
{"items": [{"aspect_id": <число>, "target": "designer" | "supplier" | "manual",
            "text": "<одно предложение: что сделать>"}]}

Правила:
1. Одно замечание — один пункт. Пиши по делу: что именно не так и что
   сделать. Максимум 20 слов на пункт.
2. НЕ цитируй регламенты, не указывай номера пунктов, статей и приложений,
   не пиши «согласно ТР ТС». Пункт читает дизайнер или менеджер, а не юрист.
3. Выбирай адресата по смыслу:
   - designer — исправляется в макете (текст, надпись, формат даты, знак,
     шрифт, опечатка, отсутствующая обязательная надпись);
   - supplier — нужны данные, которых на макете нет и которые знает только
     производитель (вид сырья, технология, происхождение, документы);
   - manual — специалист должен посмотреть сам (замеры в миллиметрах,
     графика и символы, спорная трактовка).
4. Пиши в повелительном наклонении: «Добавить…», «Запросить…», «Проверить…».
5. Без вводных слов, без «возможно», без объяснения почему — только действие
   и суть проблемы."""


def _relevant(verdicts: list[dict]) -> list[dict]:
    """Вердикты, требующие действия: нарушения и ручные проверки
    (применимые). «Соответствует» и «не применимо» действий не требуют."""
    return [v for v in verdicts
            if v.get("applicable", True)
            and v["status"] in (STATUS_VIOLATION, STATUS_MANUAL)]


def _default_target(v: dict) -> str:
    """Адресат по умолчанию для fallback и для правки человеком."""
    return "designer" if v["status"] == STATUS_VIOLATION else "manual"


def _short(text: str, limit: int = 220) -> str:
    """Последнее содержательное предложение объяснения — в промпте вердикта
    именно там стоит «что исправить / что запросить»."""
    parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", text or "") if p.strip()]
    if not parts:
        return "уточнить у специалиста"
    tail = parts[-1]
    if len(tail) < 40 and len(parts) > 1:
        tail = parts[-2] + " " + tail
    return tail[:limit]


def fallback_plan(verdicts: list[dict]) -> list[dict]:
    """План без LLM: адресат по статусу, текст — хвост объяснения."""
    return [{"aspect_id": v["id"], "aspect_name": v["name"],
             "target": _default_target(v), "text": _short(v["explanation"]),
             "source": "fallback"}
            for v in _relevant(verdicts)]


def _cache_key(model: str, payload: str) -> str:
    raw = "\x1f".join([model, SYSTEM_PROMPT, payload])
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def build_plan(report: dict, client=None, cfg: dict | None = None,
               tally: TokenTally | None = None,
               use_cache: bool = True) -> list[dict]:
    """Отчёт → список пунктов плана [{aspect_id, aspect_name, target, text}].

    client=None или сбой вызова → fallback без LLM (план всё равно есть).
    Кэш — тот же формат, что у вердиктов (config → actions.cache)."""
    cfg = cfg or load_config()
    verdicts = _relevant(report.get("verdicts", []))
    if not verdicts:
        return []
    names = {v["id"]: v["name"] for v in verdicts}
    payload = json.dumps(
        [{"aspect_id": v["id"], "аспект": v["name"], "статус": v["status"],
          "что нашли": (v.get("explanation") or "")[:900]} for v in verdicts],
        ensure_ascii=False, indent=1)

    if client is None:
        return fallback_plan(verdicts)

    acfg = cfg["actions"]
    model = os.environ[acfg["model_env"]]
    path = ROOT / acfg["cache"]
    key = _cache_key(model, payload)
    raw = None
    if use_cache:
        cached = _load_json_cache(path).get(key)
        if cached is not None:
            raw = cached["raw"]

    if raw is None:
        try:
            resp = client.chat.completions.create(
                model=model,
                response_format={"type": "json_object"},
                messages=[{"role": "system", "content": SYSTEM_PROMPT},
                          {"role": "user", "content": payload}])
            if tally is not None:
                tally.add(model, resp.usage)
            raw = json.loads(resp.choices[0].message.content)
        except Exception:  # noqa: BLE001 — план не обязан ронять отчёт
            return fallback_plan(verdicts)
        if use_cache:
            cache = _load_json_cache(path)
            cache[key] = {"model": model, "raw": raw}
            _save_json_cache(path, cache)

    items = []
    seen = set()
    for it in (raw.get("items") or []):
        if not isinstance(it, dict):
            continue
        try:
            aid = int(it.get("aspect_id"))
        except (TypeError, ValueError):
            continue
        if aid not in names or aid in seen:
            continue
        text = (it.get("text") or "").strip()
        if not text:
            continue
        target = it.get("target")
        if target not in TARGET_KEYS:
            target = _default_target(next(v for v in verdicts if v["id"] == aid))
        seen.add(aid)
        items.append({"aspect_id": aid, "aspect_name": names[aid],
                      "target": target, "text": text, "source": "llm"})

    # Аспект, который модель молча потеряла, добирается fallback'ом:
    # пропущенное замечание хуже неудачной формулировки.
    for v in verdicts:
        if v["id"] not in seen:
            items.append({"aspect_id": v["id"], "aspect_name": v["name"],
                          "target": _default_target(v),
                          "text": _short(v["explanation"]), "source": "fallback"})
    items.sort(key=lambda i: (TARGET_KEYS.index(i["target"]), i["aspect_id"]))
    return items


def apply_human_decisions(plan: list[dict], decisions: dict) -> list[dict]:
    """Правки человека поверх плана.

    decisions: {aspect_id: {"rating": "up"/"down"/None, "target": ключ или
    "none", "note": текст}}. «down» (система ошиблась) или target «none»
    убирают пункт из плана; заметка человека заменяет формулировку."""
    out = []
    for item in plan:
        d = decisions.get(item["aspect_id"]) or {}
        if d.get("rating") == "down":
            continue
        target = d.get("target") or item["target"]
        if target == "none":
            continue
        note = (d.get("note") or "").strip()
        out.append({**item, "target": target,
                    "text": note or item["text"],
                    "edited_by_human": bool(note) or target != item["target"]})
    out.sort(key=lambda i: (TARGET_KEYS.index(i["target"]), i["aspect_id"]))
    return out


def render_plan_markdown(plan: list[dict], report: dict) -> str:
    """План работ как готовый к отправке текст (без цитат регламентов)."""
    m = report.get("meta", {})
    lines = [f"# План работ по макету: {m.get('source_pdf', '—')}", ""]
    if not plan:
        lines.append("Замечаний, требующих действий, не найдено.")
        return "\n".join(lines) + "\n"
    for key, (title, _) in TARGETS.items():
        items = [i for i in plan if i["target"] == key]
        if not items:
            continue
        lines += [f"## {title}", ""]
        for i, item in enumerate(items, 1):
            lines.append(f"{i}. {item['text']} "
                         f"_(проверка: {item['aspect_name']})_")
        lines.append("")
    lines += ["---",
              "Составлено автоматически по результатам проверки макета. "
              "Инструмент предварительной проверки — финальное решение за "
              "специалистом и юристом."]
    return "\n".join(lines) + "\n"


def plan_to_docx(plan: list[dict], report: dict, path):
    """План работ в .docx (для отправки дизайнеру/поставщику как есть)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("План работ по макету", level=0)
    meta = report.get("meta", {})
    p = doc.add_paragraph(str(meta.get("source_pdf", "—")))
    p.runs[0].font.size = Pt(11)
    if not plan:
        doc.add_paragraph("Замечаний, требующих действий, не найдено.")
    for key, (title, _) in TARGETS.items():
        items = [i for i in plan if i["target"] == key]
        if not items:
            continue
        doc.add_heading(title, level=1)
        for item in items:
            para = doc.add_paragraph(item["text"], style="List Number")
            run = para.add_run(f"  (проверка: {item['aspect_name']})")
            run.italic = True
            run.font.size = Pt(9)
    doc.add_paragraph()
    tail = doc.add_paragraph(
        "Составлено автоматически по результатам проверки макета. "
        "Инструмент предварительной проверки — финальное решение за "
        "специалистом и юристом.")
    tail.runs[0].italic = True
    tail.runs[0].font.size = Pt(9)
    doc.save(str(path))
    return path
